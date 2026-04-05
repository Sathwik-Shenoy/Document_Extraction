from __future__ import annotations

import io
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, List, cast

import numpy as np
from PIL import Image


def _empty_warnings() -> List[str]:
    return []


def _empty_metadata() -> dict[str, Any]:
    return {}


@dataclass
class ExtractionResult:
    text: str
    file_type: str
    warnings: List[str] = field(default_factory=_empty_warnings)
    unreadable_tokens: int = 0
    metadata: dict[str, Any] = field(default_factory=_empty_metadata)


def get_extraction_dependency_status() -> dict[str, Any]:
    status: dict[str, Any] = {
        "pdf": {"python_package": "pdfplumber", "available": False},
        "docx": {"python_package": "python-docx", "available": False},
        "image": {
            "python_packages": {
                "opencv-python-headless": False,
                "pytesseract": False,
                "Pillow": True,
            },
            "system_binary": {"tesseract": False},
        },
    }

    try:
        import pdfplumber  # type: ignore

        status["pdf"]["available"] = True
    except Exception:
        pass

    try:
        import docx  # type: ignore

        status["docx"]["available"] = True
    except Exception:
        pass

    status["image"]["python_packages"]["opencv-python-headless"] = _safe_import_cv2() is not None
    status["image"]["python_packages"]["pytesseract"] = _safe_import_tesseract() is not None
    status["image"]["system_binary"]["tesseract"] = shutil.which("tesseract") is not None

    all_ready = (
        status["pdf"]["available"]
        and status["docx"]["available"]
        and status["image"]["python_packages"]["opencv-python-headless"]
        and status["image"]["python_packages"]["pytesseract"]
        and status["image"]["system_binary"]["tesseract"]
    )
    status["all_extractors_ready"] = all_ready
    return status


def _clean_ocr_text(text: str) -> str:
    fixes = {
        " rn ": " m ",
        "|": "I",
    }
    cleaned = f" {text} "
    for bad, good in fixes.items():
        cleaned = cleaned.replace(bad, good)
    return re.sub(r"\s+", " ", cleaned).strip()


def _safe_import_cv2():
    try:
        import cv2  # type: ignore

        return cv2
    except Exception:
        return None


def _safe_import_tesseract():
    try:
        import pytesseract  # type: ignore

        return pytesseract
    except Exception:
        return None


def _normalize_tesseract_data(data: Any) -> dict[str, Any]:
    if not isinstance(data, dict):
        return {}
    typed_data = cast(dict[str, Any], data)
    normalized: dict[str, Any] = {}
    for key, value in typed_data.items():
        normalized[str(key)] = value
    return normalized


def preprocess_image_for_ocr(raw_bytes: bytes):
    cv2: Any = _safe_import_cv2()
    if cv2 is None:
        image = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
        return np.array(image), ["opencv_not_available"]

    np_arr = np.frombuffer(raw_bytes, np.uint8)
    image: Any = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
    warnings: List[str] = []

    if image is None:
        pil_img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
        return np.array(pil_img), ["image_decode_fallback"]

    gray: Any = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    laplacian_variance: float = float(cv2.Laplacian(gray, cv2.CV_64F).var())
    if laplacian_variance < 100:
        warnings.append("blurry_image")
        sharp_kernel = np.array([[-1, -1, -1], [-1, 9, -1], [-1, -1, -1]])
        image = cv2.bilateralFilter(image, d=7, sigmaColor=50, sigmaSpace=50)
        image = cv2.filter2D(image, -1, sharp_kernel)
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    coords: Any = np.column_stack(np.where(gray < 200))
    if coords.size > 0:
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = 90 + angle
        if abs(angle) > 2:
            warnings.append("deskew_applied")
            h, w = image.shape[:2]
            matrix = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1)
            image = cv2.warpAffine(image, matrix, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    gray_arr = np.asarray(gray, dtype=np.float64)
    p10, p90 = np.percentile(gray_arr, [10, 90])
    if (p90 - p10) < 65:
        warnings.append("low_contrast")
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        gray = clahe.apply(gray)

    return gray, warnings


def extract_image_text(raw_bytes: bytes) -> ExtractionResult:
    pytesseract: Any = _safe_import_tesseract()
    image, warnings = preprocess_image_for_ocr(raw_bytes)

    if pytesseract is None:
        return ExtractionResult(
            text="",
            file_type="image",
            warnings=warnings + ["pytesseract_not_available"],
            unreadable_tokens=0,
            metadata={"ocr_confidence_mean": 0.0},
        )

    data: Any = pytesseract.image_to_data(image, config="--psm 6 --oem 3", output_type=pytesseract.Output.DICT)
    data_dict = _normalize_tesseract_data(data)
    tokens: List[str] = []
    unreadable = 0
    confidences: List[float] = []
    texts: List[Any] = list(data_dict.get("text", []))
    confs: List[Any] = list(data_dict.get("conf", []))
    for txt, conf in zip(texts, confs):
        txt = str(txt or "").strip()
        if not txt:
            continue
        conf_val = float(conf) if str(conf).replace(".", "", 1).lstrip("-").isdigit() else -1
        if conf_val >= 0:
            confidences.append(conf_val)
        if conf_val < 40:
            unreadable += 1
            tokens.append("[UNREADABLE]")
        else:
            tokens.append(txt)

    text = _clean_ocr_text(" ".join(tokens))
    avg_conf = float(np.mean(confidences) / 100.0) if confidences else 0.0
    return ExtractionResult(
        text=text,
        file_type="image",
        warnings=warnings,
        unreadable_tokens=unreadable,
        metadata={"ocr_confidence_mean": round(avg_conf, 3)},
    )


def extract_pdf_text(raw_bytes: bytes) -> ExtractionResult:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        return ExtractionResult(text="", file_type="pdf", warnings=["pdfplumber_not_available"])

    texts: List[str] = []
    table_count = 0
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            page_text = (page.extract_text(layout=True) or page.extract_text() or "").strip()
            page_chunks = [f"[PAGE {i}]\n{page_text}" if page_text else f"[PAGE {i}]"]

            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []

            for t_idx, table in enumerate(tables, start=1):
                rows: List[str] = []
                for row in table:
                    clean_cells = [((c or "").strip().replace("\n", " ")) for c in row]
                    if any(clean_cells):
                        rows.append(" | ".join(clean_cells))
                if rows:
                    table_count += 1
                    page_chunks.append(f"[TABLE {i}.{t_idx}]\n" + "\n".join(rows))

            texts.append("\n".join(page_chunks).strip())

    merged = "\n---PAGE_BREAK---\n".join(text for text in texts if text)
    return ExtractionResult(text=merged, file_type="pdf", metadata={"pages": len(texts), "tables": table_count})


def extract_docx_text(raw_bytes: bytes) -> ExtractionResult:
    try:
        import docx  # type: ignore
    except Exception:
        return ExtractionResult(text="", file_type="docx", warnings=["python_docx_not_available"])

    doc = docx.Document(io.BytesIO(raw_bytes))
    lines: List[str] = []
    heading_count = 0
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style_name = p.style.name if p.style and p.style.name else ""
        if "Heading" in style_name:
            heading_count += 1
            level_match = re.search(r"(\d+)", style_name)
            level = int(level_match.group(1)) if level_match else 1
            level = max(1, min(level, 6))
            lines.append(f"{'#' * level} {txt}")
        else:
            lines.append(txt)

    table_count = 0
    for t_idx, table in enumerate(doc.tables, start=1):
        rows: List[str] = []
        for row in table.rows:
            cells = [((cell.text or "").strip().replace("\n", " ")) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            table_count += 1
            lines.append(f"\n[TABLE {t_idx}]")
            lines.extend(rows)

    return ExtractionResult(
        text="\n".join(lines),
        file_type="docx",
        metadata={"headings": heading_count, "tables": table_count},
    )


def detect_file_type(filename: str, content_type: str | None) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}:
        return "image"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in {".txt", ".md"}:
        return "text"
    if content_type and content_type.startswith("image/"):
        return "image"
    return "text"


def extract_text(raw_bytes: bytes, filename: str, content_type: str | None) -> ExtractionResult:
    file_type = detect_file_type(filename, content_type)
    if file_type == "image":
        return extract_image_text(raw_bytes)
    if file_type == "pdf":
        return extract_pdf_text(raw_bytes)
    if file_type == "docx":
        return extract_docx_text(raw_bytes)

    text = raw_bytes.decode("utf-8", errors="replace")
    return ExtractionResult(text=text, file_type="text")
